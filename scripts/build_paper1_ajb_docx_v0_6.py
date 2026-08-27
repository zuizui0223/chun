#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Times New Roman"
BODY_PT = 12
MARGIN_IN = 1.0


def set_font(run, *, size: int = BODY_PT, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: int = BODY_PT, bold: bool | None = None) -> None:
    style.font.name = FONT
    rpr = style._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold


def add_inline(paragraph, text: str) -> None:
    """Render the limited Markdown used by the submission-clean manuscript."""
    pos = 0
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def add_continuous_line_numbering(section) -> None:
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:lnNumType"))
    if old is not None:
        sect_pr.remove(old)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def set_no_indent(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_reference_format(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    add_continuous_line_numbering(section)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    set_style_font(normal)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        set_style_font(style, bold=(style_name != "List Bullet"))
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Inches(0)

    title = doc.styles["Title"]
    set_style_font(title, size=14, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)


def build_docx(source: Path, out: Path) -> None:
    doc = Document()
    configure_document(doc)
    lines = source.read_text(encoding="utf-8").splitlines()
    section_name = "FRONT"
    title_done = False

    for raw in lines:
        line = raw.rstrip()
        if not line or line.strip() == "---":
            continue

        if line.startswith("# "):
            heading = line[2:].strip()
            if not title_done:
                paragraph = doc.add_paragraph(style="Title")
                add_inline(paragraph, heading)
                title_done = True
                section_name = "FRONT"
            else:
                paragraph = doc.add_paragraph(style="Heading 1")
                add_inline(paragraph, heading)
                section_name = heading.upper()
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, heading)
            section_name = heading.upper()
            continue

        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline(paragraph, line[4:].strip())
            continue

        if (
            line.startswith("**Running head:**")
            or line.startswith("**Authors:**")
            or line.startswith("**Affiliations:**")
            or line.startswith("**Corresponding author:**")
            or line.startswith("Manuscript received")
        ):
            paragraph = doc.add_paragraph()
            set_no_indent(paragraph)
            add_inline(paragraph, line)
            continue

        if line.startswith("**Key words:**"):
            paragraph = doc.add_paragraph()
            set_no_indent(paragraph)
            add_inline(paragraph, line)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            set_no_indent(paragraph)
            add_inline(paragraph, line[2:].strip())
            continue

        paragraph = doc.add_paragraph()
        if section_name == "LITERATURE CITED":
            set_reference_format(paragraph)
        elif section_name in {
            "FRONT",
            "ABSTRACT",
            "ACKNOWLEDGMENTS",
            "AUTHOR CONTRIBUTIONS",
            "DATA AVAILABILITY STATEMENT",
            "SUPPORTING INFORMATION",
            "FIGURE LEGENDS",
        }:
            set_no_indent(paragraph)
        add_inline(paragraph, line)

    properties = doc.core_properties
    properties.title = "Repeated flower-colour change does not imply repeated pigment-state packages in Camellia"
    properties.subject = "American Journal of Botany research article submission manuscript"
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def structural_audit(path: Path) -> dict[str, object]:
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        footer_xml = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )

    checks = {
        "continuous_line_numbers": all(
            token in document_xml
            for token in ["w:lnNumType", 'w:countBy="1"', 'w:restart="continuous"']
        ),
        "one_inch_margins": all(
            token in document_xml
            for token in [
                'w:top="1440"',
                'w:bottom="1440"',
                'w:left="1440"',
                'w:right="1440"',
            ]
        ),
        "page_number_field": " PAGE " in footer_xml,
        "double_spaced_normal_style": 'w:line="480"' in styles_xml,
        "times_new_roman_12pt_normal": all(
            token in styles_xml for token in ['w:ascii="Times New Roman"', 'w:sz w:val="24"']
        ),
    }
    failed = [name for name, passed in checks.items() if not passed]
    if failed:
        raise SystemExit(f"DOCX structural formatting audit failed: {failed}")
    return checks


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    args = parser.parse_args()

    build_docx(args.source, args.out)
    checks = structural_audit(args.out)
    if args.out.stat().st_size < 20000:
        raise SystemExit(f"DOCX output unexpectedly small: {args.out.stat().st_size} bytes")

    summary = {
        "submission_version": "v0.6",
        "source_markdown": str(args.source),
        "output_docx": str(args.out),
        "bytes": args.out.stat().st_size,
        "format": "MS Word DOCX",
        "font": "Times New Roman 12 pt body",
        "line_spacing": "double",
        "alignment": "left",
        "margins_inches": 1.0,
        "continuous_line_numbering": True,
        "sequential_page_numbering": True,
        "scientific_results_changed": False,
        "structural_checks": checks,
        "status": "AJB v0.6 DOCX built and structurally audited",
    }
    args.summary.parent.mkdir(parents=True, exist_ok=True)
    args.summary.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(summary, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
