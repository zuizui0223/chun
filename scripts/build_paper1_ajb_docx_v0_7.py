#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Times New Roman"
BODY_PT = 12
MARGIN_IN = 1.0


def set_font(run, *, size: int = BODY_PT, bold=None, italic=None) -> None:
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rpr.rFonts.set(qn(key), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: int = BODY_PT, bold=None) -> None:
    style.font.name = FONT
    rpr = style._element.get_or_add_rPr()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rpr.rFonts.set(qn(key), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold


def add_inline(paragraph, text: str) -> None:
    pos = 0
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("*"):
            set_font(paragraph.add_run(token[1:-1]), italic=True)
        else:
            set_font(paragraph.add_run(token[1:-1]))
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_continuous_line_numbering(section) -> None:
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:lnNumType"))
    if old is not None:
        sect_pr.remove(old)
    node = OxmlElement("w:lnNumType")
    node.set(qn("w:countBy"), "1")
    node.set(qn("w:distance"), "360")
    node.set(qn("w:restart"), "continuous")
    sect_pr.append(node)


def set_no_indent(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(0)
    fmt.line_spacing = 2.0
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_reference_format(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(-0.5)
    fmt.left_indent = Inches(0.5)
    fmt.line_spacing = 2.0
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = section.right_margin = Inches(MARGIN_IN)
    section.header_distance = section.footer_distance = Inches(0.5)
    add_continuous_line_numbering(section)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    set_style_font(normal)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        set_style_font(style, bold=(style_name != "List Bullet"))
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_after = style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Inches(0)
    set_style_font(doc.styles["Title"], size=14, bold=True)
    doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(source: Path, out: Path) -> None:
    doc = Document(); configure_document(doc)
    section_name = "FRONT"; title_done = False
    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line or line.strip() == "---":
            continue
        if line.startswith("# "):
            heading = line[2:].strip()
            p = doc.add_paragraph(style="Title" if not title_done else "Heading 1")
            add_inline(p, heading); title_done = True; section_name = heading.upper(); continue
        if line.startswith("## "):
            heading = line[3:].strip(); p = doc.add_paragraph(style="Heading 2"); add_inline(p, heading); section_name = heading.upper(); continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3"); add_inline(p, line[4:].strip()); continue
        if line.startswith(("**Running head:**", "**Authors:**", "**Affiliations:**", "**Corresponding author:**", "Manuscript received")):
            p = doc.add_paragraph(); set_no_indent(p); add_inline(p, line); continue
        if line.startswith("**Key words:**"):
            p = doc.add_paragraph(); set_no_indent(p); add_inline(p, line); p.add_run().add_break(WD_BREAK.PAGE); continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); set_no_indent(p); add_inline(p, line[2:].strip()); continue
        p = doc.add_paragraph()
        if section_name == "LITERATURE CITED":
            set_reference_format(p)
        elif section_name in {"FRONT", "ABSTRACT", "ACKNOWLEDGMENTS", "AUTHOR CONTRIBUTIONS", "DATA AVAILABILITY STATEMENT", "SUPPORTING INFORMATION", "FIGURE LEGENDS"}:
            set_no_indent(p)
        add_inline(p, line)

    props = doc.core_properties
    props.title = "Repeated flower-colour change does not imply repeated pigment-state packages in Camellia"
    props.subject = "American Journal of Botany research article submission manuscript"
    props.author = props.last_modified_by = props.comments = ""
    out.parent.mkdir(parents=True, exist_ok=True); doc.save(out)


def structural_audit(path: Path) -> dict[str, bool]:
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        footer_xml = "".join(archive.read(n).decode("utf-8") for n in archive.namelist() if n.startswith("word/footer") and n.endswith(".xml"))
    checks = {
        "continuous_line_numbers": all(x in document_xml for x in ["w:lnNumType", 'w:countBy="1"', 'w:restart="continuous"']),
        "one_inch_margins": all(x in document_xml for x in ['w:top="1440"', 'w:bottom="1440"', 'w:left="1440"', 'w:right="1440"']),
        "page_number_field": " PAGE " in footer_xml,
        "double_spaced_normal_style": 'w:line="480"' in styles_xml,
        "times_new_roman_12pt_normal": all(x in styles_xml for x in ['w:ascii="Times New Roman"', 'w:sz w:val="24"']),
    }
    failed = [k for k, v in checks.items() if not v]
    if failed:
        raise SystemExit(f"DOCX structural formatting audit failed: {failed}")
    return checks


def main() -> int:
    ap = argparse.ArgumentParser(); ap.add_argument("--source", type=Path, required=True); ap.add_argument("--out", type=Path, required=True); ap.add_argument("--summary", type=Path, required=True); a = ap.parse_args()
    build_docx(a.source, a.out); checks = structural_audit(a.out)
    if a.out.stat().st_size < 20000:
        raise SystemExit(f"DOCX output unexpectedly small: {a.out.stat().st_size} bytes")
    summary = {
        "submission_version": "v0.7", "source_markdown": str(a.source), "output_docx": str(a.out), "bytes": a.out.stat().st_size,
        "format": "MS Word DOCX", "font": "Times New Roman 12 pt body", "line_spacing": "double", "alignment": "left", "margins_inches": 1.0,
        "continuous_line_numbering": True, "sequential_page_numbering": True, "scientific_results_changed": False,
        "structural_checks": checks, "status": "AJB v0.7 DOCX built and structurally audited"
    }
    a.summary.parent.mkdir(parents=True, exist_ok=True); a.summary.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8"); print(json.dumps(summary, indent=2)); return 0


if __name__ == "__main__":
    raise SystemExit(main())
